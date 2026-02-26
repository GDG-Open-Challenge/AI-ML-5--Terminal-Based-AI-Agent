"""
document_loader.py — Document ingestion and vector store management (Issue #3).

Handles:
  • Loading documents from a local directory (the "upload" workflow)
  • Supported formats: PDF, DOCX, TXT, CSV
  • Chunking documents and adding them to the FAISS vector store
  • Persisting the vector store for future sessions
"""

from __future__ import annotations

import os
import csv
from typing import List

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from rich.console import Console

from config import settings

console = Console()


def _load_single_file(file_path: str) -> List[Document]:
    """Load a single file into LangChain Document objects."""
    ext = os.path.splitext(file_path)[1].lower()
    docs: List[Document] = []

    try:
        if ext == ".pdf":
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    docs.append(Document(
                        page_content=text,
                        metadata={"source": file_path, "page": i + 1},
                    ))

        elif ext == ".docx":
            from docx import Document as DocxDoc
            doc = DocxDoc(file_path)
            full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            if full_text.strip():
                docs.append(Document(
                    page_content=full_text,
                    metadata={"source": file_path},
                ))

        elif ext == ".csv":
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                reader = csv.reader(f)
                rows = list(reader)
            text = "\n".join(" | ".join(row) for row in rows)
            if text.strip():
                docs.append(Document(
                    page_content=text,
                    metadata={"source": file_path},
                ))

        else:
            # Treat as plain text
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()
            if text.strip():
                docs.append(Document(
                    page_content=text,
                    metadata={"source": file_path},
                ))

    except Exception as exc:
        console.print(f"[red]Error loading {file_path}: {exc}[/red]")

    return docs


def load_documents_from_directory(directory: str | None = None) -> List[Document]:
    """Scan a directory for supported documents and load them all.

    Args:
        directory: Path to scan.  Defaults to settings.documents_dir.

    Returns:
        List of LangChain Document objects.
    """
    directory = directory or settings.documents_dir
    directory = os.path.abspath(directory)

    if not os.path.isdir(directory):
        console.print(f"[yellow]Documents directory not found: {directory}[/yellow]")
        console.print("[yellow]Create it and place your files there to ingest them.[/yellow]")
        return []

    supported = {".pdf", ".docx", ".txt", ".csv", ".md", ".log", ".json"}
    all_docs: List[Document] = []

    for root, _dirs, files in os.walk(directory):
        for fname in sorted(files):
            ext = os.path.splitext(fname)[1].lower()
            if ext not in supported:
                continue
            full_path = os.path.join(root, fname)
            docs = _load_single_file(full_path)
            if docs:
                console.print(f"  [green]✓[/green] Loaded {fname} ({len(docs)} segment(s))")
                all_docs.extend(docs)

    console.print(f"[cyan]Total: {len(all_docs)} document segments loaded.[/cyan]")
    return all_docs


def build_vector_store(
    documents: List[Document] | None = None,
    persist_dir: str | None = None,
) -> FAISS | None:
    """Build (or rebuild) a FAISS vector store from documents.

    If *documents* is None, attempts to load from the configured documents
    directory.  The resulting store is persisted to *persist_dir*.
    """
    persist_dir = persist_dir or settings.memory_index_dir

    if documents is None:
        documents = load_documents_from_directory()

    if not documents:
        console.print("[yellow]No documents to index.[/yellow]")
        return None

    # Chunk documents
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=100,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    chunks = splitter.split_documents(documents)
    console.print(f"[cyan]Split into {len(chunks)} chunks.[/cyan]")

    # Build embeddings + store
    embeddings = HuggingFaceEmbeddings(model_name=settings.embedding_model)
    store = FAISS.from_documents(chunks, embeddings)
    store.save_local(persist_dir)
    console.print(f"[green]Vector store saved to {persist_dir}.[/green]")
    return store


def load_vector_store(persist_dir: str | None = None) -> FAISS | None:
    """Load an existing FAISS vector store from disk."""
    persist_dir = persist_dir or settings.memory_index_dir

    if not os.path.isdir(persist_dir):
        return None

    try:
        embeddings = HuggingFaceEmbeddings(model_name=settings.embedding_model)
        store = FAISS.load_local(
            persist_dir, embeddings, allow_dangerous_deserialization=True,
        )
        return store
    except Exception as exc:
        console.print(f"[red]Could not load vector store: {exc}[/red]")
        return None
